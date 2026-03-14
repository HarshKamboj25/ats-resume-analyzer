"""
utils/parser.py
===============
Handles text extraction from PDF, DOCX, and TXT files.
Libraries used:
  - pdfplumber  : extract text from PDF files
  - python-docx : extract text from DOCX files
"""

import os


def extract_text_from_file(filepath: str) -> str:
    """
    Extract plain text from a resume file.
    Supports: .pdf, .docx, .txt
    Returns the extracted text as a string.
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.pdf':
        return _extract_from_pdf(filepath)
    elif ext == '.docx':
        return _extract_from_docx(filepath)
    elif ext == '.txt':
        return _extract_from_txt(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def _extract_from_pdf(filepath: str) -> str:
    """Extract text from a PDF using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    except ImportError:
        raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")
    except Exception as e:
        raise RuntimeError(f"Failed to read PDF: {e}")


def _extract_from_docx(filepath: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"Failed to read DOCX: {e}")


def _extract_from_txt(filepath: str) -> str:
    """Read plain text file."""
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()
